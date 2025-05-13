import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import os
import re


def set_font_simsun(run):
    """设置段落字体为宋体
    
    Args:
        run: 文档中的文本运行对象
    """
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)


def convert_md_to_docx(input_md_path, output_docx_path):
    """将Markdown文件转换为Word文档
    
    Args:
        input_md_path: 输入的Markdown文件路径
        output_docx_path: 输出的Word文档路径
        
    Returns:
        str: 输出的Word文档路径
        
    Raises:
        FileNotFoundError: 当输入文件不存在时
        IOError: 当读取或写入文件出错时
    """
    try:
        # 确保输出目录存在
        output_dir = os.path.dirname(output_docx_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        # 1. 读取 Markdown 文件
        with open(input_md_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        # 2. 创建 Word 文档
        doc = Document()

        # 设置默认样式为宋体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)
        
        # 3. 逐行处理Markdown内容，识别连续行
        lines = md_content.split('\n')
        i = 0
        
        def process_bold_text(text, paragraph):
            """处理粗体文本"""
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # 粗体部分
                    bold_text = part[2:-2]
                    run = paragraph.add_run(bold_text)
                    run.bold = True
                    set_font_simsun(run)
                else:
                    # 普通文本
                    if part.strip():
                        run = paragraph.add_run(part)
                        set_font_simsun(run)
            return paragraph
        
        # 识别段落组合，例如点评中的多行属于同一个点
        def is_continuation(line, next_line):
            # 如果当前行是列表项（以- 或数字.开头）且下一行不是列表项，且下一行有内容，视为连续段落
            if (line.lstrip().startswith('-') or re.match(r'^\s*\d+\.', line.lstrip())) and next_line.strip() and not (next_line.lstrip().startswith('-') or re.match(r'^\s*\d+\.', next_line.lstrip()) or next_line.lstrip().startswith('#')):
                return True
            return False
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 检查是否需要与下一行合并（连续段落的情况）
            combined_text = line
            next_idx = i + 1
            while next_idx < len(lines) and is_continuation(line, lines[next_idx]):
                combined_text += " " + lines[next_idx].strip()
                next_idx += 1
            
            if next_idx > i + 1:
                # 有合并行，跳过已处理的行
                i = next_idx - 1
            
            # 处理标题
            if line.startswith('###'):
                text = line[3:].strip()
                heading = doc.add_heading(text, level=3)
                for run in heading.runs:
                    set_font_simsun(run)
            elif line.startswith('##'):
                text = line[2:].strip()
                heading = doc.add_heading(text, level=2)
                for run in heading.runs:
                    set_font_simsun(run)
            elif line.startswith('#'):
                text = line[1:].strip()
                heading = doc.add_heading(text, level=1)
                for run in heading.runs:
                    set_font_simsun(run)
            
            # 处理无序列表
            elif line.lstrip().startswith('-') or line.lstrip().startswith('*'):
                # 计算缩进级别
                indent = len(line) - len(line.lstrip())
                new_level = indent // 2  # 每两个空格增加一级缩进
                
                # 提取列表文本（使用组合文本如果有）
                if combined_text != line:
                    text = combined_text.lstrip()[1:].strip()
                else:
                    text = line.lstrip()[1:].strip()
                
                # 创建列表项
                para = doc.add_paragraph()
                para.style = doc.styles['List Bullet']
                
                # 设置缩进
                if new_level > 0:
                    para.paragraph_format.left_indent = Pt(18 * new_level)
                
                # 处理可能有粗体的内容
                if '**' in text:
                    process_bold_text(text, para)
                else:
                    run = para.add_run(text)
                    set_font_simsun(run)
            
            # 处理有序列表
            elif re.match(r'^\s*\d+\.', line):
                # 计算缩进级别
                indent = len(line) - len(line.lstrip())
                new_level = indent // 2  # 每两个空格增加一级缩进
                
                # 提取列表文本（使用组合文本如果有）
                if combined_text != line:
                    # 保留原始序号，去掉数字和点后的部分
                    match = re.match(r'^(\s*\d+\.)(.*)', combined_text)
                    if match:
                        prefix, text = match.groups()
                        text = text.strip()
                    else:
                        text = combined_text
                else:
                    # 保留原始序号，去掉数字和点后的部分
                    match = re.match(r'^(\s*\d+\.)(.*)', line)
                    if match:
                        prefix, text = match.groups()
                        text = text.strip()
                    else:
                        text = line
                
                # 创建列表项
                para = doc.add_paragraph()
                para.style = doc.styles['List Number']
                
                # 设置缩进
                if new_level > 0:
                    para.paragraph_format.left_indent = Pt(18 * new_level)
                
                # 处理可能有粗体的内容
                if '**' in text:
                    process_bold_text(text, para)
                else:
                    run = para.add_run(text)
                    set_font_simsun(run)
            
            # 处理完全粗体段落
            elif line.startswith('**') and line.endswith('**'):
                text = line[2:-2].strip()
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                set_font_simsun(run)
            
            # 处理包含粗体的普通段落
            elif '**' in line:
                para = doc.add_paragraph()
                process_bold_text(combined_text if combined_text != line else line, para)
            
            # 处理普通段落
            else:
                para = doc.add_paragraph(combined_text if combined_text != line else line)
                for run in para.runs:
                    set_font_simsun(run)
            
            i += 1

        # 4. 保存 Word 文档
        doc.save(output_docx_path)
        print(f"文档已保存，共 {len(doc.paragraphs)} 个段落")
        return output_docx_path
        
    except FileNotFoundError:
        raise FileNotFoundError(f"找不到输入文件: {input_md_path}")
    except Exception as e:
        import traceback
        print(f"转换过程中出错: {str(e)}")
        print(traceback.format_exc())
        raise IOError(f"转换过程中出错: {str(e)}")


# 示例用法
if __name__ == "__main__":
    input_file = "output/扫描双面_page_1_score.md"
    output_file = "output/扫描双面_page_1_score.docx"
    convert_md_to_docx(input_file, output_file)
    print(f"转换完成，文件已保存至: {output_file}")
